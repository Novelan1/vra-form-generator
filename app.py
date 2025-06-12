from dotenv import load_dotenv
import os
load_dotenv()
import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import yaml
import re
import streamlit_authenticator as stauth
from yaml.loader import SafeLoader

# === PAGE CONFIG ===
st.set_page_config(page_title="VRA Form Generator", layout="wide")

# === LOAD AUTHENTICATION CONFIG ===
with open('credentials.yaml') as file:
    config = yaml.load(file, Loader=SafeLoader)

# === SETUP AUTHENTICATOR ===
authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    os.getenv("COOKIE_SECRET"),
    config['cookie']['expiry_days']
)

# === LOGIN ===
name, authentication_status, username = authenticator.login('Login', location='main')

if authentication_status is False:
    st.error('Username or password is incorrect.')
elif authentication_status is None:
    st.warning('Please enter your username and password.')
elif authentication_status:
    authenticator.logout('Logout', 'sidebar')
    st.sidebar.success(f'Logged in as {name}')

    st.title("📄 Vulnerability Risk Assessment (VRA) Form")
    st.sidebar.header("Upload Files")

    excel_file = st.sidebar.file_uploader("Upload Excel File", type=["xlsx"])
    template_file = st.sidebar.file_uploader("Upload Word Template", type=["docx"])

    if excel_file and template_file:
        df = pd.read_excel(excel_file)
        df.columns = [col.strip().replace("\n", "_").replace(" ", "_") for col in df.columns]

        if 'AppCode' not in df.columns:
            st.error("The column 'AppCode' was not found in the Excel file.")
        else:
            appcode_list = df['AppCode'].dropna().unique().tolist()
            selected_appcode = st.selectbox("Select AppCode", appcode_list)

            if st.button("Generate Report"):
                group = df[df['AppCode'] == selected_appcode]
                if group.empty:
                    st.warning("No data found for the selected AppCode.")
                else:
                    ip_list = group['IPAddress'].dropna().unique()
                    ip_string = ", ".join(ip_list)
                    risk_levels = group['CVSSv3_RiskRating'].dropna().str.upper().unique()
                    background = f"There " + ", ".join(r.lower() for r in risk_levels) + f" vulnerabilities currently cannot be remediated in {selected_appcode}."

                    doc = Document(template_file)

                    for para in doc.paragraphs:
                        if "AppCode:" in para.text:
                            para.text = f"AppCode: \n\n{selected_appcode}"
                        elif "Systems / Servers Affected:" in para.text:
                            para.text = f"Systems / Servers Affected:\n\n{ip_string}"
                        elif "Background:" in para.text:
                            para.text = f"Background:\n\n{background}"

                    details_table = doc.tables[2]
                    existing_rows = details_table.rows[1:]
                    group_records = group.to_dict('records')

                    for i, record in enumerate(group_records):
                        if i < len(existing_rows):
                            cells = existing_rows[i].cells
                        else:
                            cells = details_table.add_row().cells

                        cells[0].text = str(i + 1)
                        cells[1].text = str(record.get('Vulnerability_Title', ''))
                        cells[2].text = str(record.get('IPAddress', ''))
                        cells[3].text = str(record.get('CVSSv3_RiskRating', ''))
                        exploit_count = str(record.get('Exploit_Count', ''))
                        cells[4].text = "NO" if exploit_count == "0" else "YES"
                        cells[5].text = str(record.get('Date_Notified', ''))
                        cells[6].text = str(record.get('SLA_Date', ''))
                        cells[7].text = str(record.get('MissedSLA', ''))

                        deviation_remark = str(record.get('DeviationRemarks', ''))
                        match = re.search(r'\b[A-Za-z]{3}\s\d{4}\b', deviation_remark)
                        cells[8].text = match.group() if match else ""

                    summary_table = doc.tables[1]
                    system_affected_col = summary_table.rows[1].cells[0].text

                    etype = df.drop_duplicates(subset=['Etype'], keep='first')[['Etype']]
                    etype_string = ', '.join(etype['Etype'].dropna().astype(str).tolist())

                    system_server_effected = df.drop_duplicates(subset=['AppCode', 'IPAddress'], keep='first')

                    def combine_ip_asset(group):
                        return '\n'.join(
                            f"{ip}\n({asset if pd.notna(asset) else '-'})"
                            for ip, asset in zip(group['IPAddress'], group['Asset_Names'])
                        )

                    system_server_effected = system_server_effected.groupby('AppCode').apply(combine_ip_asset).reset_index()
                    system_server_effected.columns = ['AppCode', 'Combined_Info']
                    ip_info = str(system_server_effected.iloc[0, 1])

                    system_affected_col = system_affected_col.replace("CTX", selected_appcode)
                    system_affected_col = system_affected_col.replace("Refer to Appendix", ip_info)
                    system_affected_col = system_affected_col.replace("Prod and DR", etype_string)

                    summary_table.cell(1, 0).text = system_affected_col

                    background_cell = summary_table.cell(1, 1)
                    background_text = background_cell.text.replace("CTX", selected_appcode).split("Justification/s:")[0]
                    background_cell.text = ""
                    para = background_cell.paragraphs[0]

                    for part in re.split(r'(\bcritical\b|\bhigh\b|\bmedium\b)', background_text, flags=re.IGNORECASE):
                        run = para.add_run(part)
                        if part.lower() in ['critical', 'high', 'medium']:
                            run.bold = True

                    severity_levels = ["CRITICAL", "HIGH", "MEDIUM", "LOW"]
                    for level in severity_levels:
                        background_cell.add_paragraph()
                        heading = background_cell.add_paragraph()
                        heading_run = heading.add_run(level)
                        heading_run.bold = True

                        level_records = group[group['CVSSv3_RiskRating'].str.upper() == level]

                        if not level_records.empty:
                            for _, row in level_records.iterrows():
                                background_cell.add_paragraph(f"\u2022 {str(row.get('Vulnerability_ID', '')).strip()}")
                                background_cell.add_paragraph(f"\u2022 {str(row.get('Vulnerability_Title', '')).strip()}")
                                background_cell.add_paragraph(f"\u2022 {str(row.get('Vulnerability_CVE_IDs', '')).strip()}")
                        else:
                            background_cell.add_paragraph("\u2022 No vulnerabilities reported.")

                    background_cell.add_paragraph()
                    background_cell.add_paragraph("Justification/s:\n<Application custodian to fill-up>")

                    report_buffer = BytesIO()
                    doc.save(report_buffer)
                    report_buffer.seek(0)

                    st.success("Report generated successfully!")
                    st.download_button(
                        label="📥 Download VRA Form",
                        data=report_buffer,
                        file_name=f"VRA_Form_{selected_appcode}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
